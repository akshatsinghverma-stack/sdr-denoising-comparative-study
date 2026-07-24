#!/usr/bin/env python3
"""
build_word_report.py — Generates report/SDR_Denoising_Project_Report.docx

A full Word-document rendition of this project: 11 sections covering both
original case studies, the severity sweep, two new case studies (time-varying
channel, 16-QAM), the boundary-aware-loss and compute-cost follow-ups, cross-
study discussion, conclusions, and limitations/future work — with a working
table of contents, page numbers, and native banded table styling.
"""

import sys
from datetime import datetime
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RESULTS_DIR = PROJECT_ROOT / "results"
FIGURES_DIR = RESULTS_DIR / "figures"
TABLES_DIR = RESULTS_DIR / "tables"
OUT_PATH = PROJECT_ROOT / "report" / "SDR_Denoising_Project_Report.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0xB0, 0x2A, 0x2A)

SNR_LEVELS = [-10, -5, 0, 5, 10, 15, 20]
METHODS_1 = ["No Processing", "LMS", "NLMS", "CNN", "Hybrid"]
METHODS_2 = ["No Processing", "LMS", "NLMS", "CNN", "Hybrid", "MMSE (Genie)"]


# ---------------------------------------------------------------------------
# Low-level helpers: TOC field, page numbers, headers/footers
# ---------------------------------------------------------------------------

def add_toc_field(doc, instr_text='TOC \\o "1-2" \\h \\z \\u',
                   placeholder_text="Right-click here and choose “Update Field” to generate the Table of Contents."):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = placeholder_text
    fld_sep.append(placeholder)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_end)


def set_update_fields_on_open(doc):
    settings = doc.settings.element
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "true")
    settings.append(uf)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_end)
    for rr in p.runs:
        rr.font.size = Pt(9)
        rr.font.color.rgb = GREY


def set_page_number_format(section, fmt="decimal", start=None):
    """Set the section's page-number display format (e.g. lowerRoman for front
    matter) and, optionally, restart numbering at `start` (used to make the
    body begin at page 1 regardless of how many front-matter pages precede it)."""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))


def add_horizontal_rule(doc, color="B02A2A", size=6):
    """A clean, plain horizontal rule (paragraph bottom border) — used on the
    title page in place of a typed line of dash characters, which renders as
    disconnected glyphs rather than a continuous line in most fonts."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), color)
    bottom.set(qn("w:space"), "1")
    border.append(bottom)
    pPr.append(border)
    return p


def add_title_header(section, title_text):
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = title_text
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = GREY
        r.italic = True


# ---------------------------------------------------------------------------
# Content helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    h.paragraph_format.keep_with_next = True
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_callout(doc, text):
    """A visually distinct callout paragraph for headline findings."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EFEFEF")
    pPr.append(shd)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), "B02A2A")
    left.set(qn("w:space"), "4")
    border.append(left)
    pPr.append(border)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, col_widths=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(font_size)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
    return table


def add_seq_caption(doc, label, text):
    """A Word-native, auto-numbered caption: 'Figure ' + SEQ field + ': ' + text.
    Uses the built-in 'Caption' style so List of Figures/Tables (TOC \\c fields)
    can find it, and so numbering recalculates automatically if order changes."""
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label} ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' SEQ {label} \\* ARABIC '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "#"
    fld_sep.append(placeholder)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_end)
    p.add_run(f": {text}")
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.color.rgb = GREY
        r.italic = True
    return p


def add_figure(doc, path, caption, width=5.8):
    if not path.exists():
        add_para(doc, f"[Figure not found: {path.name}]", italic=True, color=GREY)
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_seq_caption(doc, "Figure", caption)


def add_table_caption(doc, text):
    add_seq_caption(doc, "Table", text)


def fmt_ber(row):
    if row["ber_total_errors"] == 0:
        return f"0 err / {int(row['ber_total_bits']):,} (95% UB {row['ber_95ci_upper_bound']:.1e})"
    elif row["ber_total_errors"] < 50:
        return f"{int(row['ber_total_errors'])} err / {int(row['ber_total_bits']):,}"
    else:
        m, s = row["ber_mean"], row["ber_std"]
        return f"{m:.4f}±{s:.4f}" if m > 0.001 else f"{m:.2e}±{s:.1e}"


def case_table_rows(df, mod, snr_levels, methods, snr_col):
    rows = []
    for snr in snr_levels:
        for m in methods:
            sub = df[(df.modulation == mod) & (df.snr_db_in == snr) & (df.method == m)]
            if len(sub) == 0:
                continue
            r = sub.iloc[0]
            snr_out = f"{r[snr_col]:.2f}±{r[snr_col.replace('_mean', '_std')]:.2f}"
            rows.append([snr, m, snr_out, fmt_ber(r)])
    return rows


# ---------------------------------------------------------------------------
# Load data
# ---------------------------------------------------------------------------

df1 = pd.read_csv(TABLES_DIR / "results.csv")  # run_case1_no_isi.py's real output filename --
# "results_case1_no_isi.csv" is a differently-named legacy copy kept in sync manually; a
# code-correctness pass found this script had been silently reading that stale copy while
# report.md's hand-written tables were updated separately, risking future drift between the two.
df2 = pd.read_csv(TABLES_DIR / "results_case2_isi.csv")
df_16qam = None
try:
    df_16qam = pd.read_csv(TABLES_DIR / "results_16qam.csv")
except FileNotFoundError:
    pass

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(0.9)
section.bottom_margin = Inches(0.9)
# Front matter (title page, TOC, List of Figures/Tables) gets lowercase-roman
# page numbers, with the title page itself left unnumbered (a distinct,
# initially-empty first-page footer).
section.different_first_page_header_footer = True
add_page_number_footer(section)
set_page_number_format(section, fmt="lowerRoman")

# ===========================================================================
# Title page
# Every line below is a deliberate, explicit break — each one short enough on
# its own (well under the per-line character budget at its font size) that
# Word never needs to auto-wrap it, which is what previously left stray
# single words orphaned on their own line.
# ===========================================================================
for _ in range(2):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.line_spacing = 1.15
run = title.add_run("Comparative Study of\nDenoising Techniques\nfor SDR Signals")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = NAVY

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.line_spacing = 1.3
run = sub.add_run(
    "Classical Adaptive Filtering (LMS/NLMS)\n"
    "vs. Deep Learning (1-D CNN Autoencoder)\n"
    "vs. a Hybrid LMS+CNN Pipeline\n"
    "for BPSK / QPSK / 16-QAM Signals\n"
    "under AWGN, Multipath ISI, and Time-Varying Fading"
)
run.font.size = Pt(13)
run.italic = True
run.font.color.rgb = GREY

for _ in range(3):
    doc.add_paragraph()

author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run("Akshat Singh Verma")
run.font.size = Pt(15)
run.bold = True
run.font.color.rgb = NAVY

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
affil.paragraph_format.line_spacing = 1.3
run = affil.add_run(
    "B.Tech, Artificial Intelligence & Machine Learning\n"
    "Final Year\n"
    "University of Petroleum and Energy Studies (UPES)"
)
run.font.size = Pt(12)
run.font.color.rgb = GREY

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(datetime.now().strftime("%B %Y"))
run.font.size = Pt(11)
run.italic = True
run.font.color.rgb = GREY

doc.add_paragraph()

meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_rows = [
    ("Scope", "4 case studies + 3 connecting analyses (severity sweep, boundary-aware loss, compute cost)"),
    ("This document", "Self-contained — includes full methodology, results, verification, and conclusions"),
    ("Verification", "68-test regression suite (CI-enforced) · 8 real bugs found, root-caused, and fixed (Section 4.2, 6.1, 4.9, 3.7, 4.7, 11) · reviewed via two adversarial self-critique passes"),
    ("Source code", "Available on request — see Appendix for the experiment/module file list"),
]
for i, (k, v) in enumerate(meta_rows):
    meta_table.cell(i, 0).text = k
    meta_table.cell(i, 1).text = v
    meta_table.cell(i, 0).width = Inches(1.5)
    meta_table.cell(i, 1).width = Inches(4.5)
    for p in meta_table.cell(i, 0).paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
    for p in meta_table.cell(i, 1).paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = GREY

doc.add_page_break()

# ===========================================================================
# Abstract + Index Terms
# (Front-matter convention: an Abstract is distinct from an Executive
# Summary -- an abstract is short, technical, and IMRaD-structured for a
# reader deciding whether to read on; an executive summary is longer and
# decision-focused for a reader who may never read the rest. This report
# includes both, in that order, rather than treating them as interchangeable.)
# ===========================================================================
add_heading(doc, "Abstract", level=1)
add_para(doc,
    "This report presents a comparative study of five denoising and channel-equalization "
    "techniques — Least-Mean-Squares (LMS), Normalized LMS (NLMS), Recursive Least Squares (RLS), "
    "a 1-D convolutional neural network (CNN) autoencoder, and a Hybrid LMS→CNN cascade — against "
    "a No-Processing baseline and genie-aided linear MMSE/Zero-Forcing references, for BPSK, QPSK, "
    "and 16-QAM signals under additive white Gaussian noise (AWGN) and static and time-varying "
    "multipath fading. Four case studies each isolate a single channel variable — a memoryless AWGN "
    "channel, a static multipath (inter-symbol interference, ISI) channel, a genuinely time-varying "
    "channel, and higher-order (16-QAM) modulation — connected by three targeted follow-up analyses "
    "(an ISI-severity sweep, a boundary-aware loss-function test, and a compute-cost analysis). The "
    "central finding is that improved output SNR does not imply improved bit-error-rate (BER): on a "
    "memoryless channel, every method tested is never better than doing nothing (strictly worse in "
    "52 of 56 comparisons, tied in the remaining 4), because the raw received sample is already a "
    "sufficient statistic for detection; "
    "once real channel memory exists, the CNN reverses and can improve BER by up to 610× — though "
    "the classical adaptive filters (LMS/NLMS) remain largely unhelpful specifically for BPSK even "
    "with real ISI present — with the CNN's advantage scaling with decision-boundary crowding "
    "(BPSK < QPSK < 16-QAM), but shown to be substantially channel-specific rather than a general "
    "equalization capability once tested on channels the CNN was not trained on. Eight real "
    "implementation bugs, a missing standard baseline (Zero-Forcing), and several correlational "
    "hypotheses were found, root-caused, and — in multiple cases — directly confirmed or refuted via "
    "targeted interventions (one of which overturned an earlier, statistically well-supported but "
    "incorrect root-cause conclusion), using a discipline of checking every anomalous result against "
    "its definitional identity before accepting it. The practical recommendation is channel- and "
    "modulation-dependent: a CNN is preferred whenever real multipath structure exists, compute "
    "allows, and the deployment channel matches (or is retrained on) the training channel; RLS, once "
    "given the same numerical hardening as LMS/NLMS, is the strongest classical alternative under "
    "tight compute budgets."
)
add_para(doc,
    "Index Terms — Adaptive filtering, LMS, NLMS, RLS, convolutional neural network, denoising "
    "autoencoder, Hybrid cascade, software-defined radio, channel equalization, bit error rate, "
    "inter-symbol interference, time-varying fading, Zero-Forcing, MMSE equalizer, BPSK, QPSK, "
    "16-QAM.",
    italic=True, size=10
)

doc.add_page_break()

# ===========================================================================
# Table of Contents
# ===========================================================================
add_heading(doc, "Table of Contents", level=1)
add_toc_field(doc)
doc.add_page_break()

add_heading(doc, "List of Figures", level=1)
add_toc_field(doc, instr_text='TOC \\h \\z \\c "Figure"',
              placeholder_text="Right-click here and choose “Update Field” to generate the List of Figures.")
doc.add_page_break()

add_heading(doc, "List of Tables", level=1)
add_toc_field(doc, instr_text='TOC \\h \\z \\c "Table"',
              placeholder_text="Right-click here and choose “Update Field” to generate the List of Tables.")
doc.add_page_break()

# Now add a header (title, right-aligned) to the body pages -- this requires a
# new section so the title page / TOC don't get it.
doc.add_section(WD_SECTION.NEW_PAGE)
body_section = doc.sections[-1]
# New sections inherit "different first page" from the section they follow;
# cancel that here so the Executive Summary's own first page is numbered too.
body_section.different_first_page_header_footer = False
add_title_header(body_section, "SDR Denoising — Comparative Study Report")
add_page_number_footer(body_section)
set_page_number_format(body_section, fmt="decimal", start=1)

# ===========================================================================
# Executive Summary
# ===========================================================================
add_heading(doc, "Executive Summary", level=1)
add_para(doc,
    "This project compares five denoising/equalization approaches — LMS, NLMS, RLS, a 1-D CNN "
    "autoencoder, and a Hybrid LMS→CNN cascade — against a No-Processing baseline, for BPSK, "
    "QPSK, and 16-QAM signals from -10dB to +30dB SNR. What began as two controlled case studies "
    "grew, through follow-up questions each result raised, into four case studies plus several "
    "connecting analyses and diagnostics — every addition was triggered by a specific gap or "
    "untested claim the previous result left open, not added for its own sake."
)
add_callout(doc,
    "Headline finding: SNR improvement does not imply BER improvement, and whether any of these "
    "methods actually help depends jointly on whether the channel has structure to exploit and on "
    "how sensitive the specific modulation's decision geometry is to the specific distortion that "
    "structure introduces."
)
add_para(doc, "Key results, in the order they were established:")
add_numbered(doc, [
    "Case Study 1 (memoryless AWGN, no ISI): every method — LMS, NLMS, CNN, Hybrid — is NEVER "
    "BETTER than doing nothing at all, despite large apparent SNR gains, because the raw received "
    "sample is already the Bayes-optimal decision variable when the channel has no memory (strictly "
    "worse in 52/56 comparisons, tied in the remaining 4 — CNN specifically, at 15/20dB, after its "
    "reconstruction-bug fix eliminated its floor down to No-Processing's already-zero error count).",
    "Case Study 2 (real multipath ISI): the result can reverse — QPSK shows equalization winning "
    "by up to ~470x — but BPSK barely benefits and classical filters are frequently worse than "
    "doing nothing even with ISI present. Getting this result trustworthy required finding and "
    "fixing three real bugs (a receiver gain bug, an SNR-calibration bug, and a channel-causality "
    "bug), each caught by checking a value against what it must equal by definition.",
    "An ISI severity sweep turns the two-case comparison into an actual crossover curve: QPSK at "
    "10dB goes from \"equalization hurts\" at zero ISI to \"equalization helps by 5.6x\" at full "
    "severity.",
    "A genie-aided linear MMSE equalizer was built as a reference bound — an eighth real bug in its "
    "autocorrelation computation initially made it look like a ceiling CNN could exceed, but once "
    "fixed, it ties or beats CNN at high SNR, exactly as a genie with perfect channel knowledge should.",
    "Case Study 3 (time-varying channel) tested this project's most load-bearing design decision "
    "— freezing LMS/NLMS after the preamble — and found a fourth bug (the decision-directed safety "
    "gate is dead code at SPS>1) while also refuting the general principle behind that decision: "
    "on a genuinely time-varying channel, LMS tracking earns back a real, substantial BER win.",
    "Case Study 4 (16-QAM) confirmed the decision-boundary-crowding hypothesis more dramatically "
    "than BPSK/QPSK alone predicted, revealing a hard, un-closeable BER floor for No-Processing "
    "not seen for BPSK/QPSK. (An earlier version of this finding also claimed the genie MMSE "
    "ceiling turns into an outright regression here -- that was this project's eighth real bug, "
    "corrected: the properly-computed genie MMSE bound stays close to optimal at every modulation "
    "tested, not a regression at any of them.)",
    "A boundary-aware CNN loss confirmed \"MSE is boundary-blind\" as a real, partial, "
    "modulation-dependent cause of Case Study 1's headline result, closing up to 78.5% of QPSK's "
    "excess BER — but not BPSK's.",
    "A compute-cost analysis found CNN needs ~106-212x more arithmetic per sample than LMS/NLMS, "
    "yet is paradoxically faster in wall-clock terms on this project's CPU stack — a finding about "
    "implementation efficiency, not algorithmic cost, with a concrete hardware-deployment answer "
    "either way.",
    "The LMS-vs-NLMS decision-directed asymmetry from Case Study 3 was root-caused with an actual "
    "intervention: NLMS's instantaneous-power step-size normalization measurably inflates its "
    "effective step size during channel fades; flooring that normalization measurably closes the "
    "gap in every tested cell and flips two of four from \"DD doesn't help\" to \"DD helps.\"",
    "The high-SNR CNN error floor was first attributed, on strong correlational grounds (81.1% "
    "same-position overlap between loss functions, χ²=168.0 p=3.5×10⁻³⁶ clustering by window "
    "phase), to a window/overlap-add reconstruction artifact — a conclusion a later "
    "code-correctness critique pass found to be WRONG: the real cause was a bug (this project's "
    "seventh) that zero-filled a signal's uncovered tail samples. Fixing it collapsed BPSK's floor "
    "and QPSK's floor at 15-20dB almost entirely, leaving only a small, genuine residual at QPSK "
    "10dB now explained by high-noise-magnitude samples.",
    "An outside-reviewer critique pass found every CNN result trains and tests on the exact same "
    "known channel, unlike LMS/NLMS/MMSE/ZF which all re-estimate or are handed the channel "
    "per-trial. Testing this directly confirmed a real generalization gap: a CNN frozen on one "
    "channel degrades up to 129.71x on a held-out channel (QPSK, 10dB) where LMS/NLMS degrade only "
    "1.17-1.41x — revising, not overturning, the project's practical recommendation.",
    "Training the CNN once (rather than across several seeds) showed no evidence of a problem in "
    "one narrow, tested regime: training-draw variance was smaller than test-time Monte Carlo "
    "variance in all 4 conditions tested (Case Study 1's CNN only, 2 SNR points) — not yet a "
    "project-wide validation.",
    "RLS, added as a third classical baseline after fixing a real safety bug (an unguarded "
    "decision-directed default): a first pass without tail-averaging found RLS worse than LMS/NLMS "
    "at low-to-moderate SNR, but adding tail-averaging (matching LMS/NLMS's own hardening) and "
    "re-running showed RLS now matches or beats both at every SNR level tested — the original "
    "theoretical motivation confirmed, once RLS is given a fair, equally-hardened comparison.",
    "Zero-Forcing (ZF), a standard equalizer baseline this project was missing, was added after "
    "searching comparable published work rather than re-reading this project's own code — once "
    "the eighth bug above was fixed, it confirmed the textbook prediction directly: the slightly "
    "worse of the two genie equalizers for QPSK, both comfortably beating No-Processing.",
])

doc.add_page_break()

# ===========================================================================
# How to Read This Report: Pipeline & Terminology
# ===========================================================================
add_heading(doc, "How to Read This Report: Pipeline & Terminology", level=1)
add_para(doc,
    "Every result in this document comes from the same basic experiment, run under different "
    "channel conditions and with different modulations. One pass through it looks like this:"
)
add_numbered(doc, [
    "Generate a long random bit sequence and modulate it (map bits to symbols) using one of "
    "BPSK, QPSK, or 16-QAM.",
    "Optionally pulse-shape and oversample the symbol stream (Case Studies 2-4), then pass it "
    "through a simulated radio channel: additive noise (AWGN), and — in most case studies — "
    "multipath distortion (ISI) that smears each symbol into its neighbors, sometimes also "
    "changing over time (Case Study 3).",
    "Apply one of five processing methods to the received, corrupted signal: do nothing "
    "(the baseline), LMS or NLMS adaptive filtering, a CNN autoencoder, a Hybrid LMS→CNN cascade, "
    "or (Case Studies 2/4 only) a genie-aided MMSE equalizer that is handed the exact channel for "
    "reference.",
    "Demodulate the processed signal back into bits and compare against the original bit sequence "
    "to compute a bit error rate (BER), alongside a signal-to-noise ratio (SNR) measurement of the "
    "processed signal itself.",
    "Repeat many times (a Monte Carlo trial) with independently regenerated bits and noise at each "
    "SNR level, and report the mean and standard deviation across trials — never a single run.",
])
add_para(doc,
    "Every method except \"do nothing\" and MMSE first sees a short preamble — a known reference "
    "sequence transmitted before the real data — which it uses to train/adapt before being "
    "evaluated on the rest of the signal."
)

add_heading(doc, "Glossary", level=2)
glossary = [
    ("BPSK / QPSK / 16-QAM", "Digital modulation schemes carrying 1, 2, or 4 bits per transmitted symbol, respectively — more bits per symbol means more, closer-together constellation points, and less noise margin per bit."),
    ("SNR", "Signal-to-Noise Ratio, in dB — how strong the desired signal is relative to background noise."),
    ("BER", "Bit Error Rate — the fraction of transmitted bits recovered incorrectly. The metric that actually matters for a working communication link."),
    ("AWGN", "Additive White Gaussian Noise — the standard baseline noise model: random, uncorrelated, Gaussian-distributed interference added to the signal."),
    ("ISI", "Inter-Symbol Interference — distortion where one symbol's energy leaks into neighboring symbols, typically from multipath (a signal arriving via more than one path/delay)."),
    ("LMS / NLMS", "Least Mean Squares / Normalized LMS — classical adaptive filters that iteratively adjust a set of weights (\"taps\") to undo channel distortion, learning from the known preamble."),
    ("CNN", "A 1-D convolutional neural network, here used as a denoising autoencoder: trained to map a noisy/distorted signal window back to a clean one."),
    ("Hybrid", "This project's LMS-then-CNN cascade: LMS does a coarse first pass, and a CNN is trained to clean up what LMS leaves behind."),
    ("MMSE (Genie)", "Minimum Mean Square Error linear equalizer, given the exact channel and noise level as a reference upper bound — no real adaptive method has this information, hence \"genie.\""),
    ("RRC / SPS", "Root-Raised-Cosine pulse shaping and Samples-Per-Symbol — standard techniques for band-limiting a transmitted signal, used from Case Study 2 onward (SPS=4)."),
    ("Decision-Directed (DD)", "An adaptive filter mode that keeps adapting after the preamble, using its own (possibly wrong) decisions as a substitute reference — powerful but riskier than freezing the filter, as Section 3.2 and 6.2 discuss in depth."),
    ("Monte Carlo trial", "One complete, independent repetition of an experiment with freshly generated random bits and noise — used to report a mean and spread rather than a single, potentially lucky or unlucky, run."),
    ("Rule of three", "A standard statistical convention: when zero errors are observed in N trials, the true error rate is not reported as a bare \"0\" but bounded above (95% confidence) by approximately 3/N."),
    ("MAC / FLOP", "Multiply-Accumulate operation / Floating-Point Operation — units of computational cost, used in Section 8 to compare methods by actual arithmetic cost rather than just runtime."),
]
add_table(doc, ["Term", "Meaning"], glossary, col_widths=[1.7, 4.3], font_size=9)

doc.add_page_break()

# ===========================================================================
# 1. Objective
# ===========================================================================
add_heading(doc, "1. Objective", level=1)
add_para(doc,
    "Compare five denoising/equalization approaches against a No-Processing baseline, for BPSK, "
    "QPSK, and 16-QAM signals corrupted by AWGN, multipath ISI, and (in one follow-up) time-varying "
    "fading. The project is organized as four case studies plus several connecting analyses and "
    "diagnostics, each generated directly by a gap or unverified claim in a previous result:"
)
add_bullets(doc, [
    "Case Study 1 (Section 3): a memoryless AWGN channel (no pulse shaping, no ISI).",
    "Case Study 2 (Section 4): the same methods, unchanged, on a channel with real ISI (RRC pulse "
    "shaping at 4x oversampling + a static multipath channel), plus a genie-aided MMSE equalizer "
    "and, as a follow-up, RLS as a third classical baseline (Section 4.9).",
    "ISI Severity Sweep (Section 5): connects Case Studies 1 and 2 by scaling channel severity "
    "continuously between them.",
    "Case Study 3 (Section 6): a genuinely time-varying channel, testing whether freezing LMS/NLMS "
    "after the preamble is still the right call once the channel actually moves, and root-causing "
    "why LMS benefits from tracking while NLMS does not.",
    "Case Study 4 (Section 7): 16-QAM, testing whether decision-boundary crowding keeps sharpening "
    "the equalization story as the constellation gets more crowded.",
    "Boundary-Aware Loss and Training-Variance (Sections 3.7-3.8): turns Case Study 1's diagnosis "
    "of why CNN loses to No-Processing into a tested, partially-confirmed causal claim, root-causes "
    "the resulting high-SNR error floor, and quantifies how much training randomness (vs. test-time "
    "randomness) contributes to reported BER variance.",
    "Compute-Cost Analysis (Section 8): supplements every BER/SNR comparison with a "
    "hardware-portable answer to how much each method actually costs to run.",
])

# ===========================================================================
# 2. Methodology
# ===========================================================================
add_heading(doc, "2. Shared Methodology", level=1)
add_bullets(doc, [
    "BPSK / QPSK: 100,000 symbols per trial (reduced to 20,000-30,000 for the newer, "
    "diagnostic-scoped follow-ups — severity sweep, Case Studies 3/4, boundary-loss — consistent "
    "with the severity sweep's original precedent).",
    "AWGN: complex, correctly power-scaled, SNR sweep [-10, -5, 0, 5, 10, 15, 20] dB for the two "
    "main case studies (shifted to [0,10,20,25,30]dB for 16-QAM, which needs far more SNR headroom "
    "before ISI effects are even visible).",
    "Monte Carlo evaluation: 10 independent trials per (modulation, SNR, method) in the two main "
    "case studies (5 in the diagnostic follow-ups), each with an independently regenerated bit "
    "sequence and noise realization. Results are reported as mean ± std, never a single point "
    "estimate, and BER is tracked as raw (error_count, bits_tested) pairs so an exact-zero BER is "
    "reported honestly via the rule-of-three 95% upper bound rather than a misleading bare 0.0.",
    "LMS / NLMS: 16 taps, μ=0.01 / μ=0.5 (ε=1e-6), 1000-symbol preamble using the true clean signal "
    "as reference, then FROZEN for the remainder of the signal by default (Section 3.2; revisited "
    "in Section 6).",
    "CNN: a small 1-D Conv autoencoder (3 encoder + 3 decoder layers; 6,890 parameters, verified — "
    "see Section 8 — correcting this project's earlier \"~15k\" estimate), window=128, stride=64, "
    "trained across all SNR levels jointly, Adam 1e-3, ≤50 epochs with EarlyStopping.",
    "Hybrid: LMS (coarse) → CNN retrained on LMS's residual output (fine).",
    "MMSE (Genie, Case Study 2/4 only): closed-form linear MMSE equalizer given the exact channel "
    "and noise level — an upper-bound reference, not a competing adaptive method.",
    "No Processing: raw noisy signal straight to hard-decision demod — the falsifiability baseline.",
])

doc.add_page_break()

# ===========================================================================
# 3. Case Study 1
# ===========================================================================
add_heading(doc, "3. Case Study 1: Memoryless AWGN Channel (No ISI)", level=1)

add_heading(doc, "3.1 Configuration", level=2)
add_para(doc,
    "No pulse shaping, no oversampling (SPS=1), no multipath — each received sample corresponds "
    "to exactly one symbol with zero statistical dependence on any other sample."
)

add_heading(doc, "3.2 Fixing LMS/NLMS Divergence", level=2)
add_para(doc,
    "The initial implementation diverged catastrophically at low input SNR (BPSK -10dB: LMS output "
    "SNR of -15.07dB, worse than the -10dB input). Two fixes were required:"
)
add_numbered(doc, [
    "The preamble training itself was unstable, not just the decision-directed phase afterward — "
    "a single realization run for ~1000 iterations with noise-dominated input showed error power "
    "spiking >1000x within the classical stability bound. Clamping to 20% of the bound (rather than "
    "90%) fixed this, and the frozen weight vector is the average of the last ~70% of the preamble "
    "trajectory (Polyak/Ruppert averaging) rather than the raw final iterate.",
    "Decision-directed continuation was tried and rejected as unsafe — engaging-and-diverging in "
    "~30-40% of random trials at borderline SNR. Monte Carlo testing across many seeds showed "
    "permanently freezing after the preamble matches or beats decision-directed continuation at "
    "every tested SNR level, so it is off by default (revisited in Section 6, where this "
    "justification is refined).",
])

add_heading(doc, "3.3 Results — BPSK", level=2)
rows = case_table_rows(df1, "BPSK", SNR_LEVELS, METHODS_1, "snr_db_out_mean")
add_table_caption(doc, "SNR and BER results, BPSK, Case Study 1 (memoryless AWGN)")
add_table(doc, ["SNR in (dB)", "Method", "SNR out (mean±std, dB)", "BER (mean±std or exact)"], rows)

doc.add_page_break()
add_heading(doc, "3.4 Results — QPSK", level=2)
rows = case_table_rows(df1, "QPSK", SNR_LEVELS, METHODS_1, "snr_db_out_mean")
add_table_caption(doc, "SNR and BER results, QPSK, Case Study 1 (memoryless AWGN)")
add_table(doc, ["SNR in (dB)", "Method", "SNR out (mean±std, dB)", "BER (mean±std or exact)"], rows)

doc.add_page_break()
add_heading(doc, "3.5 Figures", level=2)
c1fig = FIGURES_DIR / "case1_no_isi"
add_figure(doc, c1fig / "ber_vs_snr_BPSK.png", "BER vs Input SNR, BPSK, Case Study 1")
add_figure(doc, c1fig / "ber_vs_snr_QPSK.png", "BER vs Input SNR, QPSK, Case Study 1")

doc.add_page_break()
add_heading(doc, "3.6 Key Finding: SNR Improvement Does Not Imply BER Improvement", level=2)
add_callout(doc,
    "Every single method, at every SNR level, for both modulations, is never better than doing "
    "nothing at all — strictly worse in 52/56 comparisons, tied (both zero errors) in the "
    "remaining 4 — despite CNN posting up to +18.92dB of apparent SNR gain. Verified exhaustively, "
    "not sampled. The 4 ties are CNN at BPSK/QPSK's two highest SNRs, after Section 3.7's "
    "reconstruction-bug fix collapsed its floor there to exactly 0."
)
add_para(doc,
    "Under a memoryless AWGN model, the raw received sample is already a sufficient statistic for "
    "detecting its symbol, and hard-sign detection on it is the Bayes-optimal (minimum-BER) "
    "receiver — there is no BER headroom above doing nothing. Every method pools information across "
    "multiple samples (16 taps or a 128-sample window), the majority of which belong to "
    "statistically independent, unrelated symbols. Verified concretely: the converged LMS filter at "
    "20dB places 98.1% of its weight on the current-sample tap (matching the theoretical Wiener "
    "shrinkage factor SNR/(1+SNR)=0.990), while the other 15 taps, though ~320x smaller, are not "
    "exactly zero — finite-sample training noise leaves them holding energy from unrelated symbols, "
    "pure interference for the current bit decision. The CNN compounds this: MSE loss has no "
    "explicit penalty for crossing the decision boundary, so it trades sign flips for a larger "
    "reduction in squared error elsewhere."
)
add_para(doc,
    "This result is specific to a memoryless, ISI-free channel — it directly motivated Case Study "
    "2.", italic=True
)

add_heading(doc, "3.7 Follow-Up: Is “MSE Is Boundary-Blind” a Tested Causal Claim?", level=2)
add_para(doc,
    "Section 3.6’s diagnosis was an explanation, not a tested claim. A boundary-hinge loss (MSE + "
    "a margin penalty for crossing the decision boundary, added to an architecturally-identical "
    "CNN) was built and evaluated on Case Study 1's exact configuration to test it directly."
)
add_callout(doc,
    "Confirmed for QPSK at low-to-moderate SNR: the boundary-hinge loss closes 60-79% of the CNN's "
    "excess BER over No-Processing (78.5% at -10dB exactly). Not confirmed for BPSK: hinge made it "
    "slightly worse. A third, loss-independent error floor appears at 10-20dB for both modulations, "
    "unaffected by either loss function."
)
add_para(doc,
    "A bit-level flip diagnostic confirms this is mechanistic: at QPSK -10dB, MSE's harmful/"
    "beneficial flip counts are 14,683/12,611 (net harm 2,072); Hinge's are 6,037/5,592 (net harm "
    "445) — both counts roughly halve under the boundary-aware loss, exactly the predicted effect. "
    "The predicted trade-off is also confirmed: Hinge has systematically lower output SNR than MSE "
    "(e.g. QPSK 15dB: 18.04dB vs 21.28dB) — it genuinely trades accuracy for margin, and that only "
    "pays off where boundary ambiguity, not some other structural error source, dominates. Full "
    "detail: report/findings_boundary_aware_loss.md."
)
add_callout(doc,
    "CORRECTED FINDING. This floor was first attributed, on strong correlational grounds (81.1% "
    "same-position overlap between MSE/Hinge, χ²=168.0 p=3.5×10⁻³⁶ non-uniform phase clustering, "
    "only 15.1% overlap with No-Processing), to a window/overlap-add reconstruction artifact. A "
    "later code-correctness critique pass found the real cause: a bug, not a reconstruction "
    "property — this project's SEVENTH real bug."
)
add_para(doc,
    "`reconstruct_from_windows` (`src/cnn_autoencoder.py`) silently left any trailing samples not "
    "covered by a full window at their initialized value of exactly 0.0 — for this diagnostic's "
    "configuration, the last 40 samples of every processed signal, every trial. A hard-decision "
    "demodulator reads a raw 0.0 as a fixed, deterministic bit (BPSK's real>=0 rule always decides "
    "bit 1) regardless of what was transmitted — a constant, SNR-independent error source hiding "
    "exactly in the 10-20dB range where a \"floor\" would be visible. Fixed by giving "
    "reconstruct_from_windows a fallback array (the noisy input) so uncovered positions pass "
    "through unmodified instead of zero — matching the convention src/mmse_equalizer.py already "
    "used. Five new regression tests (tests/test_cnn_autoencoder.py) lock in both the gap and the fix."
)
add_callout(doc,
    "Re-running the exact diagnostic with the fix: BPSK's floor (a constant 83/100,000 errors at "
    "10, 15, AND 20dB) collapses to 1/100,000 and 0/100,000. QPSK's floor at 15-20dB (161/200,000) "
    "collapses completely to 0. The phase-clustering statistic recomputes to χ²≈0.74 (p≈0.86, "
    "indistinguishable from uniform) on the corrected error population — the original χ²=168.0 "
    "result was itself manufactured by the bug: the always-wrong zero-filled tail sits at a fixed, "
    "narrow range of window-phase values, fabricating spurious clustering out of deterministic "
    "tail errors."
)
add_para(doc,
    "One genuine, smaller residual floor survives the fix, at QPSK 10dB only (365→218/200,000, a "
    "real ~40% reduction, not full elimination). Re-examining just this population now points to "
    "the OTHER original hypothesis: 43.7% of remaining error positions are also wrong for "
    "No-Processing (up from 15.1%), and noise magnitude at these positions averages 2.83x typical "
    "(up from 1.39-1.64x) — consistent with genuinely hard noise realizations, not a positional "
    "artifact. This also retroactively explains the triangular-weighting intervention's negative "
    "result below: it could never have worked, since reweighting cannot restore a sample no window "
    "covers at all. Full corrected detail: report/findings_cnn_high_snr_floor.md."
)
add_callout(doc,
    "The named intervention was tested before this correction, and produced a NEGATIVE result at "
    "the time, reported honestly. A triangular overlap-add weighting (downweighting each window's "
    "own edges) did NOT shrink the floor: BPSK error counts were identical at every SNR tested; "
    "QPSK was unchanged or very slightly WORSE at 10dB. As the correction above explains, this was "
    "the correct result for a different reason than originally understood."
)

add_heading(doc, "3.8 Follow-Up: Is Training the CNN Once Enough?", level=2)
add_para(doc,
    "This project's practice (every case study) is to train each CNN/Hybrid model exactly once per "
    "modulation, then report Monte Carlo variance across test-time noise/bit-sequence draws only. The "
    "variance contributed by the training draw itself was previously unmeasured."
)
add_callout(doc,
    "In one narrow, tested regime (Case Study 1's CNN, 2 SNR points), 5 independently-seeded CNNs "
    "evaluated on identical test data show no evidence of training-draw variance dominating "
    "test-time variance (ratio 0.00-0.32x across all 4 conditions tested) — not yet a project-wide "
    "validation, and the estimate itself rests on only 4-5 samples per condition."
)
add_para(doc,
    "At 15dB, five independently-trained BPSK models produced the exact same mean BER to 6 decimal "
    "places; QPSK's five were within 3% of each other. At 0dB (noisier), the ratio is largest (0.32x "
    "for QPSK) but still well under 1. This is also consistent with the high-SNR floor being a "
    "structural, architecture-driven property (Section 3.7) rather than something training randomness "
    "could shift. Full detail: report/findings_training_variance.md."
)

doc.add_page_break()

# ===========================================================================
# 4. Case Study 2
# ===========================================================================
add_heading(doc, "4. Case Study 2: RRC Pulse-Shaped + Multipath ISI Channel", level=1)

add_heading(doc, "4.1 Configuration", level=2)
add_para(doc,
    "Identical to Case Study 1 with exactly one change: 4x oversampling with RRC pulse shaping, "
    "plus a static 3-tap multipath channel h=[1.0, 0.4+0.3j, -0.1+0.1j]."
)

add_heading(doc, "4.2 Three Normalization Bugs Found and Fixed", level=2)
add_para(doc,
    "Getting the SNR metric — and the channel itself — trustworthy took finding and fixing three "
    "real bugs, all caught by checking a value against what it is definitionally required to equal:"
)
add_numbered(doc, [
    "Amplitude/gain bug in receiver_frontend: reusing the transmit RRC filter as the receive "
    "matched filter without renormalizing double-counted its scaling — recovered symbols at "
    "amplitude ±4.0 instead of ±1.0. Fixed by dividing by the filter's own energy.",
    "SNR-calibration bug: the matched filter gives the signal a coherent-combining gain noise "
    "doesn't get, plus a further channel-specific gain — together a consistent +5.77dB offset in "
    "No-Processing's measured SNR. Fixed via empirical calibration (inject a known probe SNR, "
    "measure what's actually observed, use the gap as calibration) — verified to within 0.001-"
    "0.03dB of nominal.",
    "Causal-convolution bug in add_multipath, found while building the ISI severity sweep (Section "
    "5) — its severity=0 anchor should exactly reproduce Case Study 1's result and instead showed "
    "CNN winning by up to 1000x. Root cause: add_multipath used mode=\"same\" convolution (correct "
    "for a symmetric filter, wrong for a causal channel model) — it silently shifted the signal by "
    "~1 sample, verified with h=[1,0,0] (must be identity) instead returning a shifted signal. "
    "Fixed with causal convolution; three regression tests added.",
])
add_para(doc,
    "All three fixes changed the actual signal/noise injected, so the entire sweep was re-run "
    "after each. The tables below reflect the fully corrected run, now including a genie-aided "
    "MMSE equalizer (Section 4.7)."
)

doc.add_page_break()
add_heading(doc, "4.3 Results — BPSK", level=2)
rows = case_table_rows(df2, "BPSK", SNR_LEVELS, METHODS_2, "snr_db_out_vs_postisi_ref_mean")
add_table_caption(doc, "SNR and BER results, BPSK, Case Study 2 (RRC pulse shaping + multipath ISI)")
add_table(doc, ["SNR in (dB)", "Method", "SNR out vs post-ISI ref (dB)", "BER (mean±std or exact)"], rows)

doc.add_page_break()
add_heading(doc, "4.4 Results — QPSK", level=2)
rows = case_table_rows(df2, "QPSK", SNR_LEVELS, METHODS_2, "snr_db_out_vs_postisi_ref_mean")
add_table_caption(doc, "SNR and BER results, QPSK, Case Study 2 (RRC pulse shaping + multipath ISI)")
add_table(doc, ["SNR in (dB)", "Method", "SNR out vs post-ISI ref (dB)", "BER (mean±std or exact)"], rows)

doc.add_page_break()
add_heading(doc, "4.5 Verification Before Trusting These Numbers", level=2)
add_bullets(doc, [
    "No train/test leakage: CNN/Hybrid training completes entirely before any test-trial signal is "
    "generated; all 10 test-trial bit sequences confirmed bit-for-bit distinct from training and "
    "each other.",
    "CNN's \"0 errors\" verified via held-out generalization (a freshly-trained CNN on 5 never-"
    "used seeds: 0 errors in 500,000 bits, matching the in-pipeline result) and decision-margin "
    "analysis (minimum margin 0.49-0.78, comfortably clear of the boundary).",
    "The naive plausibility check (Q-function at nominal SNR) was wrong and revealing: actual "
    "measured SNR is only ~7-9dB, and Q-function at that level predicts hundreds of errors — "
    "resolved by recognizing the Q-function assumes Gaussian noise, while this residual is "
    "structured/non-Gaussian, confirmed directly via the margin analysis.",
    "High-SNR LMS/NLMS/Hybrid \"regressions\" against No-Processing (Fisher's exact test, 9 cells) "
    "were re-checked with an actual multiple-comparisons correction, not just a caveat: Bonferroni "
    "(family-wise, corrected α=0.0056) leaves 3/9 significant; Benjamini-Hochberg FDR (the more "
    "defensible choice for this exploratory family) leaves 7/9 significant — every cell except "
    "NLMS's two 20dB points (p=0.063 each). Under BH: LMS/Hybrid's regressions are real; NLMS's is "
    "not statistically distinguishable from chance at either 20dB point — a small, genuine "
    "misadjustment noise floor for LMS/Hybrid, not a meaningful practical difference either way.",
])

add_heading(doc, "4.6 Key Finding: Equalization Helps — But How Much Depends on Modulation", level=2)
add_callout(doc,
    "BPSK barely benefits from equalization and LMS/NLMS/Hybrid are frequently WORSE than doing "
    "nothing. QPSK is the opposite: every method routs No-Processing by 1-3 orders of magnitude."
)
add_para(doc,
    "BPSK's single binary decision boundary is forgiving enough that this channel's ISI mostly "
    "wasn't hurting it to begin with, so there's little room for a classical filter to win, while "
    "that filter's own estimation noise is a real, constant cost. QPSK's four-quadrant boundary is "
    "far more sensitive to the same channel's phase/amplitude smearing — there is real structure to "
    "exploit, and it shows. “Does equalization help” depends not just on “is there "
    "ISI,” but on how sensitive the modulation's decision geometry is to the specific "
    "distortion ISI introduces."
)

add_heading(doc, "4.7 The Genie-Aided MMSE Bound", level=2)
add_para(doc,
    "A closed-form linear MMSE equalizer given the channel's exact taps and noise level (a genie "
    "no adaptive method has)."
)
add_callout(doc,
    "CORRECTED FINDING (this project's EIGHTH real bug). The original analysis below concluded "
    "QPSK genie MMSE gets systematically worse at high SNR due to a real linear-equalizer noise/ISI "
    "tradeoff. That conclusion was wrong. design_mmse_equalizer (src/mmse_equalizer.py) computed "
    "its autocorrelation via np.correlate(g, g), which silently returns the CONJUGATE of the value "
    "its own documented Wiener-Hopf formula requires -- negligible for BPSK, severe for QPSK/16-QAM. "
    "Fixed via return np.conj(c_full[idx]), verified against a from-scratch ground-truth solve."
)
add_para(doc,
    "Re-running Case Study 2 at full scale with the fix reverses the finding completely. QPSK genie "
    "MMSE BER at 10/15/20dB was 0.0193/0.0041/4.9e-4 (38,505/8,312/980 errors of 2,000,000) -- worse "
    "than every adaptive method, including No-Processing at 20dB (0 errors). AFTER THE FIX: "
    "1,923/0/0 errors -- MMSE now BEATS every adaptive method at 10dB (CNN: 2,789; LMS: 6,786; "
    "NLMS: 7,272; No-Processing: 18,692) and TIES CNN and No-Processing at 0 from 15dB onward, "
    "exactly what a genie with perfect channel and noise knowledge should do. The phase-bias/"
    "tap-count/frequency-response investigation below was a real, honest attempt to verify an "
    "anomalous result -- the discipline was right, it just never inspected the one place (the "
    "autocorrelation's own algebra) where the bug actually lived. This also reverses the original "
    "\"CNN exceeds the linear-equalizer ceiling\" conclusion: with the bug fixed, CNN ties MMSE at "
    "high SNR and loses to it at 10dB -- the genie bound is strong, not one CNN casually surpasses."
)
add_para(doc,
    "Original (now-superseded) analysis, preserved for transparency: at low-to-moderate SNR MMSE "
    "behaved as expected, a real edge over LMS/NLMS. At high SNR for QPSK, the buggy output "
    "appeared to have zero mean phase error (unbiased) but a real residual phase spread (~16° std) "
    "large enough to cross QPSK's ±45° boundary while harmless for BPSK's ±180°; increasing taps "
    "from 16 to 51 changed zero errors, ruling out \"too short a filter\" -- both checks were "
    "individually sound, neither happened to inspect the actual bug. This channel's frequency "
    "response was also computed directly (FFT of h=[1.0, 0.4+0.3j, -0.1+0.1j]): |H(f)| dips to 0.50 "
    "near f≈-0.46 cycles/symbol, a 10.1dB drop from the peak — a real attenuation region, though "
    "this mechanism is no longer the explanation for the (corrected) finding above."
)

doc.add_page_break()
add_heading(doc, "4.8 Figures", level=2)
c2fig = FIGURES_DIR / "case2_isi"
add_figure(doc, c2fig / "ber_vs_snr_BPSK.png", "BER vs Input SNR, BPSK, Case Study 2")
add_figure(doc, c2fig / "ber_vs_snr_QPSK.png", "BER vs Input SNR, QPSK, Case Study 2")
add_figure(doc, c2fig / "constellation_BPSK.png", "Received constellation, BPSK, Case Study 2")
add_figure(doc, c2fig / "constellation_QPSK.png", "Received constellation, QPSK, Case Study 2")
add_figure(doc, c2fig / "convergence_BPSK.png", "Adaptive filter convergence, BPSK, Case Study 2")
add_figure(doc, c2fig / "convergence_QPSK.png", "Adaptive filter convergence, QPSK, Case Study 2")

doc.add_page_break()
add_heading(doc, "4.9 Follow-Up: RLS as a Third Classical Baseline", level=2)
add_para(doc,
    "RLS was originally excluded from Case Study 2 to keep exactly one variable (the channel) "
    "different from Case Study 1. Before this first real use, src/rls_filter.py — implemented but "
    "unused — needed a real safety fix (this project's FIFTH real bug): it defaulted to unguarded "
    "decision-directed continuation, the same broken |d-y|-threshold pattern already fixed for "
    "LMS/NLMS (Section 3.2). Hardened to frozen-by-default first (16 new regression tests), then "
    "run at Case Study 2's full scale."
)
add_para(doc,
    "First pass (RLS without tail-averaging): not the simple \"RLS wins\" result predicted — RLS "
    "was measurably WORSE than LMS/NLMS at low-to-moderate SNR (5-19% more errors) and only "
    "strictly better at high SNR (zero errors at 15-20dB where LMS/NLMS retain a residual floor)."
)
add_callout(doc,
    "Corrected finding: adding Polyak/Ruppert tail-averaging (the same hardening LMS/NLMS already "
    "had) and re-running produces a real, substantial improvement at EVERY SNR level — RLS now "
    "matches or beats LMS/NLMS across the entire tested range, not just at high SNR. A first attempt "
    "wrongly concluded \"zero effect\" by comparing against a stale results file — caught and "
    "corrected by checking the file's modification timestamp against the run's actual completion."
)
rls_rows = [
    ["BPSK -10/-5/0", "405,908 / 301,200 / 136,225", "355,427 / 250,848 / 113,564", "1.14-1.20x", "1.01-1.07x"],
    ["BPSK 5/10", "25,202 / 198", "16,611 / 81", "1.52x / 2.44x", "1.41x / 2.96x"],
    ["BPSK 15-20", "0 / 0", "0 / 0", "-", "inf"],
    ["QPSK -10/-5/0", "867,057 / 706,102 / 436,240", "788,298 / 632,816 / 380,709", "1.10-1.15x", "1.01-1.03x"],
    ["QPSK 5/10", "143,117 / 8,093", "113,821 / 4,371", "1.26x / 1.85x", "1.11x / 1.55x"],
    ["QPSK 15-20", "0 / 0", "0 / 0", "-", "inf"],
]
add_table_caption(doc, "RLS error counts before/after tail-averaging, and improvement vs. LMS (after)")
add_table(doc, ["Condition", "RLS errors (before)", "RLS errors (after)", "Improvement", "vs. LMS (after)"], rls_rows)
doc.add_paragraph()
add_para(doc,
    "Checked with the same formal significance test Section 4.5 applies elsewhere, not eyeballed: "
    "Fisher's exact test on RLS-before-vs-after is significant at every SNR level (p<10⁻¹² "
    "throughout); RLS(after)-vs-LMS/NLMS(after) is significant from -10 to 10dB and consistent with "
    "a tie at 15-20dB where counts near zero (e.g. BPSK 20dB: RLS 0 vs. NLMS 4 errors, p=0.125 — "
    "reported honestly as not distinguishable, not claimed as a further win). \"vs. LMS(after)\" is "
    "≥1.0 at every SNR level tested — RLS with tail-averaging is never worse than LMS/NLMS anywhere "
    "in this study's range, confirming the original future-work rationale once RLS is given the "
    "same hardening LMS/NLMS already had. RLS is also cheaper per-sample than NLMS in this "
    "implementation despite its higher algorithmic complexity, another instance of Section 8's "
    "compute-vs-wall-clock finding. Full detail, including the process lesson about verifying "
    "re-run results against fresh file timestamps: report/findings_rls_comparison.md."
)
rlsfig = FIGURES_DIR / "rls_comparison"
add_figure(doc, rlsfig / "ber_vs_snr_BPSK.png", "BER vs Input SNR incl. RLS, BPSK")
add_figure(doc, rlsfig / "ber_vs_snr_QPSK.png", "BER vs Input SNR incl. RLS, QPSK")

doc.add_page_break()
add_heading(doc, "4.10 Follow-Up: Zero-Forcing (ZF) — a Missing Baseline Found in Published Work", level=2)
add_para(doc,
    "A self-critique pass deliberately searched comparable published equalization literature "
    "(rather than re-reading this project's own code) and found Zero-Forcing (ZF) is a standard "
    "textbook linear-equalizer baseline this project was missing entirely — it had genie-aided MMSE "
    "(Section 4.7) but no ZF, despite the two normally being presented side by side. design_zf_"
    "equalizer (src/mmse_equalizer.py) is a two-line addition reusing existing machinery — exactly "
    "design_mmse_equalizer with the noise term dropped — verified to converge to MMSE as its noise "
    "term shrinks, and to recover symbols with zero errors at negligible noise."
)
add_callout(doc,
    "RE-RUN after fixing Section 4.7's autocorrelation-conjugate bug (this project's eighth), since "
    "this comparison uses the same equalizer code and was affected identically. BPSK is unchanged "
    "(tracks MMSE closely, both reach zero errors at 15-20dB). For QPSK, 20dB now ties all three "
    "methods at 0 errors, so the clearest comparison is at 10dB: MMSE has 1,923/2,000,000 errors, "
    "ZF has 1,942/2,000,000 -- ZF still (very slightly) the worse of the two genie equalizers, "
    "exactly the textbook prediction, just at a realistic magnitude (previously: MMSE 38,582, ZF "
    "45,373 at the same SNR -- both worse than No-Processing's 18,692, which a genie should never "
    "be). Both genies now comfortably beat No-Processing rather than losing to it."
)
add_para(doc,
    "This is not a new, independent mechanism — it is the same linear-equalizer noise-enhancement "
    "tradeoff from Section 4.7, now shown to scale monotonically with how noise-blind the equalizer "
    "is (ZF > MMSE > adaptive/learned methods, in order of how much this cost is paid). Full detail: "
    "report/findings_zf_comparison.md."
)

doc.add_page_break()
add_heading(doc, "4.11 Follow-Up: Does the CNN Generalize Across Channels, or Memorize the One It Was Trained On?", level=2)
add_para(doc,
    "An outside-reviewer critique pass — deliberately unfamiliar with this project's internal "
    "history, briefed to find the single biggest weakness a viva examiner would raise — flagged a "
    "gap no prior self-critique round had named: every CNN result in this report trains and "
    "evaluates on the exact same known channel (MULTIPATH in run_case2_isi.py, reused verbatim for "
    "training and every test trial). LMS/NLMS re-estimate the channel from the preamble every trial; "
    "genie MMSE/ZF are handed the true channel every trial. Only the CNN implicitly assumes "
    "train-time and test-time channels match — untested, despite feeding directly into Section 9.5's "
    "central recommendation."
)
add_para(doc,
    "One CNN is trained once per modulation on the matched (Case Study 2) channel, then evaluated — "
    "with NO retraining — on five held-out channels differing in delay spread, tap count, and tap "
    "phase (not just a severity-scaled version of the same taps, which Section 5 already covers and "
    "always trains/tests at matched severity). LMS, NLMS, and genie MMSE are run on the identical "
    "held-out channels as a like-for-like reference, since they adapt/measure regardless of which "
    "channel is presented (experiments/test_cnn_channel_generalization.py, reduced scale)."
)
add_callout(doc,
    "CONFIRMED: the CNN's advantage is substantially channel-specific, not a general equalization "
    "capability. At QPSK 10dB, a channel with sign-flipped tap phases degrades CNN BER by 129.71x "
    "relative to its own matched-channel performance (1.65e-3 -> 2.14e-1), while LMS and NLMS on "
    "the IDENTICAL held-out channel degrade only 1.17x and 1.41x. Across all five held-out channels "
    "at this SNR, CNN degradation ranges 1.22x-129.71x versus LMS/NLMS's tight 0.67x-1.41x band."
)
add_para(doc,
    "The held-out channel is objectively somewhat harder for every method (even No-Processing "
    "degrades 6.19x there), but catastrophically so only for the frozen CNN. The effect is smaller "
    "at BPSK (worst-case 2.08x vs. LMS/NLMS's 0.98-1.10x), consistent with BPSK's wide margin being "
    "less sensitive to mechanisms this project keeps finding QPSK sensitive to (Sections 3.7, 4.6, "
    "4.7, 7), and vanishes at 20dB where every method reaches near-zero BER on every channel."
)
add_callout(doc,
    "This does not overturn Case Study 2's headline finding -- the CNN still substantially "
    "outperforms LMS/NLMS whenever train and test channels match, the condition every prior case "
    "study actually tested. But that advantage has been measured under an unrealistic best case: a "
    "receiver that already knows, implicitly, exactly which channel it will face. Section 9.5's "
    "recommendation is revised accordingly. Full detail: report/findings_cnn_channel_generalization.md."
)

doc.add_page_break()

# ===========================================================================
# 5. Severity Sweep
# ===========================================================================
add_heading(doc, "5. ISI Severity Sweep: Locating the Crossover", level=1)
add_para(doc,
    "Case Study 1 (no ISI) and Case Study 2 (one specific channel) gave opposite answers — but "
    "that's only two data points. This sweep scales the channel's tap magnitudes by a severity "
    "factor from 0 to 1 and re-runs LMS, CNN, and the genie MMSE at each level."
)
sev_rows = [
    ["BPSK 0dB", "LMS", "0.71x", "0.70x", "0.70x", "0.71x", "0.72x"],
    ["BPSK 0dB", "CNN", "0.93x", "0.95x", "0.98x", "1.01x", "1.05x"],
    ["QPSK 0dB", "LMS", "0.85x", "0.86x", "0.87x", "0.89x", "0.91x"],
    ["QPSK 0dB", "CNN", "0.96x", "0.97x", "0.99x", "1.04x", "1.07x"],
    ["QPSK 10dB", "LMS", "0.52x", "0.66x", "1.09x", "1.82x", "2.68x"],
    ["QPSK 10dB", "CNN", "0.64x", "0.80x", "1.48x", "3.15x", "5.61x"],
]
add_table_caption(doc, "BER improvement ratio vs. ISI severity")
add_table(doc, ["Case", "Method", "sev=0.00", "0.25", "0.50", "0.75", "1.00"], sev_rows)
doc.add_paragraph()
add_callout(doc,
    "QPSK at 10dB: a monotonic crossover from “equalization hurts” through breakeven "
    "(~severity 0.5) to “equalization helps substantially” (up to 5.61x) — a genuine "
    "dose-response curve, not two disconnected snapshots."
)

doc.add_page_break()
add_heading(doc, "Figures", level=2)
sevfig = FIGURES_DIR / "severity_sweep"
add_figure(doc, sevfig / "severity_crossover_BPSK.png", "BER improvement ratio vs ISI severity, BPSK")
add_figure(doc, sevfig / "severity_crossover_QPSK.png", "BER improvement ratio vs ISI severity, QPSK")

doc.add_page_break()

# ===========================================================================
# 6. Case Study 3: Time-Varying Channel
# ===========================================================================
add_heading(doc, "6. Case Study 3: Time-Varying Channel", level=1)
add_para(doc,
    "Every stability decision through Section 5 was justified by “nothing to track on a "
    "static channel, so continuation is pure downside risk.” That claim was never tested, "
    "because every channel used so far is time-invariant. This section builds a genuinely "
    "time-varying channel (drifting multipath taps, sinusoidal or random-walk models) and tests it."
)
add_callout(doc,
    "Correction (found via a later adversarial self-critique pass): this channel model was "
    "originally described as drifting slowly enough to \"look static across the 1000-symbol "
    "preamble.\" That claim was asserted, not measured, and is FALSE for the flat-fading config "
    "actually used in Section 6.4's headline result — measured directly, it swings by 60.5-90.7% "
    "of its base tap magnitude within the preamble itself, across the 5 seeds tested."
)
add_para(doc,
    "This does not invalidate Section 6.4's Frozen-vs-DD BER comparison — both methods see the "
    "identical channel/noise realization in every trial, so their difference is real regardless — "
    "but it means the channel is continuously time-varying from the start of the signal, not "
    "static-then-drifting as originally described, and the \"Frozen\" baseline is itself a fit to "
    "an already-moving channel. Full detail: report/findings_preamble_drift_correction.md."
)

add_heading(doc, "6.1 Finding 1 — a Fourth Bug: the DD Reliability Gate Is Dead Code at SPS>1", level=2)
add_para(doc,
    "The safety gate deciding whether decision-directed (DD) continuation is trustworthy compares "
    "hard decisions on the preamble tail against the exact clean reference — built and validated "
    "only for SPS=1, where every clean sample IS exactly a constellation point. At SPS=4, the clean "
    "reference is a continuously-varying RRC waveform matching a constellation point at only 1-in-4 "
    "samples, so the mismatch check structurally reads ~75-80% “mismatch” regardless of "
    "SNR or channel behavior, and the gate never opens: measured 0/9 engagement across every "
    "combination tested. Invisible until now because no script in this project had ever actually "
    "set enable_decision_directed=True."
)

add_heading(doc, "6.2 Finding 2 — Bypassing the Gate Is Catastrophic, and It's Not About Time-Variation", level=2)
add_para(doc,
    "Forcing DD to run anyway makes it catastrophic — BER up to ~10,000x worse than frozen weights "
    "across -5 to 15dB, both modulations, both drift modes. Root-caused rather than reported at "
    "face value: a weight-vector inspection shows DD converges to an entirely different, delay-"
    "shifted filter, and — the decisive check — running the identical test on Case Study 2's "
    "ORIGINAL STATIC channel reproduces the same catastrophic failure. The failure has nothing to "
    "do with time-variation: DD's per-sample decide-and-adapt loop is only correct when every "
    "sample is a symbol (SPS=1); at SPS=4, three of four samples are transitional pulse-shape "
    "values, and forcing decisions on them corrupts the filter regardless of channel behavior."
)

add_heading(doc, "6.3 Finding 3 — the Clean Test: SPS=1, Genuine Flat Fading, Gate Untouched", level=2)
add_para(doc,
    "A control at SPS=1 (where the gate works correctly) with a genuinely time-varying flat-fading "
    "channel: the gate behaves exactly as designed (0% engagement at low SNR, rising to 20-100% as "
    "SNR increases)."
)
add_callout(doc,
    "LMS decision-directed tracking earns back its keep: pooled BER drops 29-40% vs. Frozen at "
    "10-15dB, both modulations, beating even No-Processing. NLMS does NOT show the same benefit — "
    "flat-to-worse at every SNR, with a 37.5% engage-and-diverge rate matching this project's "
    "originally-documented “~30-40% of trials” almost exactly."
)
add_para(doc,
    "This asymmetry has since been root-caused with an actual intervention, not left as "
    "correlational evidence (report/findings_lms_nlms_asymmetry.md): an instrumented "
    "reimplementation of both filters' update equations — verified to reproduce production BER "
    "exactly before being trusted — measured the correlation between the channel's instantaneous "
    "envelope and each method's effective step size across all 31 DD-engaged trials. The result "
    "splits perfectly by method, zero exceptions: all 16 engaged LMS trials show a fade shrinking "
    "its update-vector norm; all 15 engaged NLMS trials show the opposite — a fade inflating NLMS's "
    "effective step size, exactly when wrong decisions are most likely."
)
add_callout(doc,
    "The named falsification test was then actually performed: a new min_norm_power parameter "
    "(default 0.0, zero effect on any other experiment) floors NLMS's normalization denominator. "
    "Re-running the flat-fading control with the floor set to 25% of each trial's own measured "
    "mean windowed power MEASURABLY CLOSES NLMS's disadvantage in every tested POOLED cell — the "
    "Frozen/DD improvement ratio rises from 0.77x to 0.98x (BPSK 10dB), 0.72x to 0.83x (BPSK "
    "15dB), 0.96x to 1.14x (QPSK 10dB), and 1.12x to 1.80x (QPSK 15dB) — QPSK's two cells FLIP "
    "from \"DD doesn't help\" to \"DD measurably helps.\""
)
add_para(doc,
    "The pooled averages hide real per-trial variance, reported rather than smoothed over: of 50 "
    "individual trials, 27 changed measurably under the floor — 16 improved, 11 regressed. Ten "
    "regressions are negligible (<0.001 BER), but one is not: BPSK 15dB trial 1 went from 0.229 to "
    "0.429 under flooring, close to random-guessing. The floor is a net improvement on average, not "
    "a guaranteed improvement every time — a production deployment would need a less aggressive or "
    "smoothed floor to manage that per-trial risk. A 10%/25%/50% floor-fraction sensitivity check "
    "confirms the conclusion isn't fragile to the 25% value chosen: 10% is clearly too small "
    "(worse than 25% in every tested cell), while 50% performs comparably or better — the pattern "
    "a real mechanism should show, not flat insensitivity. BPSK's gap to LMS's reported benefit is reduced "
    "but not fully closed — the normalization mechanism is a real, causally-confirmed contributor, "
    "but not the sole one for BPSK specifically. This turns a best-supported explanation into a "
    "demonstrated cause: NLMS's instantaneous normalization is not merely inferred but shown, by "
    "direct intervention, to produce the destabilizing feedback loop that LMS's fixed step size "
    "structurally cannot have."
)

add_heading(doc, "6.4 Conclusion", level=2)
add_para(doc,
    "Freezing remains the right default today — but primarily because DD's implementation needs a "
    "symbol-rate-aware fix before SPS>1 channels can use it at all, not because time-variation "
    "would never make continuation worthwhile. The broader stated justification (“nothing to "
    "track, pure downside risk”) is refuted as a general principle: on a genuinely time-"
    "varying channel, LMS tracking is a real, substantial win at adequate SNR. The original "
    "engage-and-diverge concern is reproduced, not eliminated. Full detail: "
    "report/findings_timevarying_channel.md."
)

doc.add_page_break()
add_heading(doc, "Figures", level=2)
tvfig = FIGURES_DIR / "timevarying"
add_figure(doc, tvfig / "frozen_vs_dd_BPSK.png", "Frozen vs. Decision-Directed, BPSK, time-varying channel")
add_figure(doc, tvfig / "frozen_vs_dd_QPSK.png", "Frozen vs. Decision-Directed, QPSK, time-varying channel")
add_figure(doc, tvfig / "lms_nlms_asymmetry_diagnostic.png",
           "LMS vs. NLMS update-norm and channel envelope, illustrative diverging trial")

doc.add_page_break()

# ===========================================================================
# 7. Case Study 4: 16-QAM
# ===========================================================================
add_heading(doc, "7. Case Study 4: Higher-Order Modulation (16-QAM)", level=1)
add_para(doc,
    "Case Study 2 found BPSK and QPSK experience the same ISI channel very differently, attributed "
    "to decision-boundary crowding. 16-QAM (16 tightly-packed points) tests whether this keeps "
    "sharpening. Implementation self-tested with a zero-noise round-trip (0 bit errors over 20,000 "
    "bits) before use — this project's standard practice, which caught a real Gray-mapping bug in a "
    "bonus 8-PSK implementation before it could contaminate any result."
)

if df_16qam is not None:
    rows16 = case_table_rows(df_16qam, "16QAM", [0, 10, 20, 25, 30], METHODS_2,
                              "snr_db_out_vs_postisi_ref_mean")
    add_table_caption(doc, "SNR and BER results, 16-QAM, Case Study 4")
    add_table(doc, ["SNR in (dB)", "Method", "SNR out vs post-ISI ref (dB)", "BER (mean±std or exact)"], rows16)
    doc.add_page_break()

add_heading(doc, "BER Improvement Ratio Comparison", level=2)
mod_compare = [
    ["LMS", "never wins (≤0.7x)", "31.5x (15dB)", "break-even at 0dB, 480x at 25dB"],
    ["NLMS", "never wins", "43.0x (15dB)", "break-even at 0dB, 610x at 25dB"],
    ["CNN", "1.8x (10dB, best)", "473x (15dB)", "557x at 25dB"],
]
add_table_caption(doc, "BER improvement ratio vs. No-Processing, by modulation")
add_table(doc, ["Method", "BPSK best (§4.6)", "QPSK best (§4.6)", "16-QAM (this study)"], mod_compare)

doc.add_paragraph()
add_callout(doc,
    "Confirmed, more dramatically than expected: equalization wins by up to 610x for 16-QAM, "
    "versus QPSK's 473x. But the mechanism is richer than just “a bigger ratio.”"
)
add_bullets(doc, [
    "No-Processing hits a genuinely new, hard BER floor (0.0755-0.0814 from 20-30dB) — never seen "
    "for BPSK or QPSK — deterministic ISI alone is enough to permanently cross 16-QAM's tight "
    "margins regardless of noise level.",
    "CORRECTED: \"the genie MMSE linear equalizer becomes actively WORSE than doing nothing\" was "
    "this project's eighth real bug (Section 4.7's autocorrelation-conjugate error), not a real "
    "effect. Re-run with the fix: MMSE errors at 20/25/30dB are 46/46/43 out of 500,000 bits -- "
    "essentially tied with CNN (30/21/21) and ~880-970x BETTER than No-Processing, not worse.",
])
add_para(doc,
    "Full detail: report/findings_higher_order_modulation.md.", italic=True
)

doc.add_page_break()
add_heading(doc, "Figures", level=2)
qamfig = FIGURES_DIR / "higher_order_modulation"
add_figure(doc, qamfig / "ber_vs_snr_16QAM.png", "BER vs Input SNR, 16-QAM")
add_figure(doc, qamfig / "constellation_16QAM.png", "Received constellation, 16-QAM")
add_figure(doc, qamfig / "convergence_16QAM.png", "Adaptive filter convergence, 16-QAM")

doc.add_page_break()

# ===========================================================================
# 8. Compute Cost
# ===========================================================================
add_heading(doc, "8. Compute-Cost Analysis: When Is CNN Actually Worth It?", level=1)
add_para(doc,
    "Every comparison above is BER/SNR-only, plus a runtime_sec_mean column that measures this "
    "specific software implementation's wall-clock speed — not a hardware-portable answer to how "
    "much compute each method fundamentally needs."
)

compute_rows = [
    ["LMS", "16 taps", "64.0", "64.0", "1.0×"],
    ["NLMS", "16 taps", "64.0", "64.0", "1.0× (+1 div/sample)"],
    ["MMSE (Genie)", "16 taps", "64.0", "64.0", "1.0× (+ one-time design solve)"],
    ["CNN", "6,890 params", "6,784.0", "13,568.0", "106.0× (up to 212×)"],
    ["Hybrid", "6,906", "6,848.0", "13,632.0", "107.0×"],
]
add_table_caption(doc, "Compute cost by method: state size, MACs/sample, and relative cost vs. LMS")
add_table(doc, ["Method", "State size", "Real MACs/sample (window-norm.)",
                "Real MACs/sample (stride-honest)", "Relative MAC cost vs LMS"], compute_rows)

doc.add_paragraph()
add_callout(doc,
    "CNN needs ~106-212x more raw arithmetic per sample than LMS/NLMS/MMSE, yet is FASTER in "
    "wall-clock terms on this project's CPU — because LMS/NLMS are pure per-sample Python loops "
    "(interpreter overhead dwarfs the actual 64 MACs of useful work), while CNN inference is "
    "vectorized, BLAS-backed TensorFlow. This is an implementation-efficiency finding, not an "
    "algorithmic-complexity one — dedicated hardware (FPGA/DSP/MCU) would see the 106-212x MAC gap "
    "reassert itself as a real throughput requirement."
)
add_para(doc,
    "Concrete decision rule: 64-640 MMACs/s (LMS/NLMS/MMSE) fits comfortably on a basic Cortex-M4/"
    "M7-class microcontroller; 6.8-136 GMACs/s (CNN) does not — it needs an applications processor, "
    "DSP, or NPU (roughly 20-500x beyond MCU-class throughput). On a bare microcontroller, or with "
    "BPSK/no-ISI, CNN is not worth deploying. On a platform with multi-GMACs/s-to-TOPS-class "
    "headroom, and a channel with real, modulation-sensitive ISI structure (QPSK or 16-QAM here), "
    "CNN's compute multiplier buys a large, measured BER improvement that is a good trade for most "
    "platforms above the microcontroller tier. Full derivation: report/compute_cost.md."
)

doc.add_page_break()

# ===========================================================================
# 9. Cross-Case-Study Discussion
# ===========================================================================
add_heading(doc, "9. Cross-Case-Study Discussion", level=1)
add_para(doc,
    "Case Study 1 and Case Study 2 differ in exactly one variable (the channel), and reach "
    "opposite conclusions about whether denoising/equalization helps BER. Neither result is "
    "“the answer” in isolation — together they establish that the value of adaptive "
    "filtering or learned denoising for BER is conditional on the channel actually having structure "
    "to exploit. Case Study 4 shows this isn't a two-point BPSK/QPSK curiosity but a continuing "
    "trend: 16-QAM's tighter margins push the effect further in the same direction on both ends."
)
add_para(doc,
    "The SNR-vs-BER disconnect recurs via two different, independently-verified mechanisms: "
    "boundary-blind MSE training (Case Study 1) and structured non-Gaussian residual dragging down "
    "an average SNR metric (Case Study 2, CNN). (A third variant was originally reported here -- "
    "phase-structured residual error making the genie MMSE bound look disconnected from BER -- but "
    "that was this project's eighth real bug, not a genuine example; see Section 4.7's correction.) "
    "One repeated lesson: an aggregate SNR or MSE number should not be trusted as a BER proxy "
    "without checking the actual decision-margin distribution."
)
add_para(doc,
    "Case Study 3 revises how this project talks about its most load-bearing design decision: "
    "freezing LMS/NLMS after the preamble is still right, but the justification on record was "
    "broader than what was actually verified, and testing it surfaced a fourth bug that had been "
    "masking a separate, real incompatibility. “We tested it and it works” claims need to "
    "specify exactly what was tested."
)
add_para(doc,
    "Taken together, this project's four case studies plus three connecting analyses form a "
    "coherent structure in which each new question was generated directly by a gap in a previous "
    "result, rather than chosen in advance — arguably the more defensible way to scope an open-"
    "ended comparative study."
)

doc.add_page_break()

# ===========================================================================
# 9.5 Practical Recommendation
# ===========================================================================
add_heading(doc, "9.5 Practical Recommendation: Which Method Should You Actually Use?", level=1)
add_para(doc,
    "Every finding above says \"it depends.\" That is the scientifically honest answer, but a "
    "comparative study should still be able to say, concretely, what to actually deploy in a given "
    "situation — each recommendation below cites the specific evidence behind it."
)
rec_rows = [
    ["Clean/near-memoryless channel, any modulation", "No-Processing (do nothing)",
     "Case Study 1: no method ever beats doing nothing -- strictly worse in 52/56 comparisons, tied in the remaining 4 (§3.6)."],
    ["Real ISI, BPSK", "CNN if compute allows; else No-Processing",
     "§4.6: BPSK barely benefits from any equalizer; only CNN shows a small, consistent edge (1.1-1.8x)."],
    ["Real ISI, QPSK or 16-QAM", "CNN",
     "§4.6/7: CNN wins by 100-610x — the single largest, most consistent effect in this project."],
    ["ISI channel, microcontroller-class compute only", "LMS or NLMS, moderate-high SNR",
     "§8: 64-640 MMACs/s vs. CNN's 6.8-136 GMACs/s; still a real 1-3 order-of-magnitude win for QPSK."],
    ["ISI channel, classical filters, any SNR", "RLS (tail-averaged)",
     "§4.9: once hardened like LMS/NLMS, RLS matches or beats both at every SNR level tested."],
    ["Genuinely time-varying channel", "LMS-DD (fix SPS>1 bug first); NLMS-DD only with the floor fix",
     "§6.4: LMS-DD beats Frozen by 29-40%; NLMS-DD needs the normalization floor to do the same."],
    ["Closed-form reference, channel knowable", "MMSE over ZF",
     "§4.10: ZF (no noise-awareness) is consistently the slightly worse of the two genie equalizers for QPSK; both comfortably beat No-Processing."],
    ["Deployment channel may drift from the training channel", "Retrain the CNN on the deployment channel, or budget for degraded BER",
     "§4.11: a CNN frozen on one channel degrades up to 129.71x on a held-out channel (QPSK, 10dB) vs. 1.17-1.41x for LMS/NLMS on the same channel."],
]
add_table_caption(doc, "Practical decision matrix: situation, recommendation, and supporting evidence")
add_table(doc, ["Situation", "Recommended method", "Why (evidence)"], rec_rows, font_size=8)
doc.add_paragraph()
add_callout(doc,
    "If forced to name one single default: a CNN denoising autoencoder, deployed only when the "
    "channel has real multipath/ISI structure, the deployment channel is expected to match (or be "
    "retrained on) the training channel, and the platform has at least applications-processor-class "
    "compute. It never loses to No-Processing once real channel structure exists, and its advantage "
    "grows, not shrinks, as the situation gets more realistic -- but that advantage is measurably "
    "channel-specific (§4.11), not a general equalization capability."
)
add_para(doc,
    "How this compares to published work: literature on classical adaptive equalizers [1], [3], [4] "
    "reports RLS beating LMS by ~45-51% BER reduction (~1.8-2.0x) for BPSK/QPSK/8PSK; a separate "
    "ZF/LMS/RLS comparative study [2] motivated adding the Zero-Forcing baseline in Section 4.10. "
    "This project's first-pass RLS result looked like a contradiction (RLS losing to a heavily-"
    "hardened LMS/NLMS) — but that was an unfair comparison against an unhardened RLS. Once RLS was "
    "given equivalent hardening (§4.9), the corrected result aligns well with the published range "
    "(1.1x-3.8x, bracketing ~1.8-2.0x in the middle). Separately, published comparisons of deep "
    "denoising autoencoders against classical filters [5], [6] report neural approaches pulling "
    "ahead specifically when channel characteristics are complex or non-linear — a pattern this "
    "project's own Case Study 1 vs. Case Study 2/4 comparison reproduces independently, from first "
    "principles. Full citations: References."
)

doc.add_page_break()

# ===========================================================================
# 10. Conclusions
# ===========================================================================
add_heading(doc, "10. Conclusions", level=1)
add_numbered(doc, [
    "SNR improvement does not imply BER improvement, and the relationship between them is "
    "channel-dependent, not universal — the headline finding of this project. Its corollary: even "
    "within one channel regime, an aggregate SNR/MSE metric can misrepresent BER for structural "
    "reasons distinct from the cross-channel finding.",
    "Classical adaptive filters (LMS/NLMS) are prone to real, verifiable instabilities that require "
    "deliberate diagnosis and fixing, not just parameter tuning.",
    "CNN-based denoising outperforms classical filtering once channel structure exists to exploit, "
    "but can be actively harmful on a structureless channel — the same architecture, opposite "
    "outcomes, purely as a function of the channel.",
    "The Hybrid cascade's underperformance in Case Study 1 was mostly an artifact of LMS's "
    "divergence bug; once fixed, Hybrid is competitive with standalone CNN in Case Study 2.",
    "Rigorous small-sample-size handling (Monte Carlo averaging, rule-of-three bounds, significance "
    "testing) materially changed which findings survived scrutiny and should be standard practice.",
    "Definitional identities are a cheap, powerful bug-finder — all three Case Study 2 bugs, plus "
    "the fourth found in Case Study 3, were caught by knowing what a value must equal by definition.",
    "A regression test suite (tests/, 68 tests) now codifies every invariant found this way.",
    "A design decision's stated justification can be narrower, or broader, than what was actually "
    "verified — only testing the untested case reveals which (Case Study 3).",
    "A mechanistic explanation is only as strong as the experiment that tests it — the boundary-"
    "aware-loss follow-up turned one explanatory sentence into three separately-evidenced sub-"
    "claims.",
    "Decision-boundary crowding is a continuum, not a BPSK/QPSK-specific curiosity — 16-QAM is the "
    "sharpest point on it tested so far, both for equalization's benefit and for No-Processing's "
    "un-closeable error floor (the genie MMSE \"ceiling cost\" once suspected here was this "
    "project's eighth bug, not a real effect — see Section 4.7/7's correction).",
    "Compute cost and BER cost do not move together, and both need to be reported — CNN's 106-212x "
    "compute multiplier is invisible in this project's own wall-clock numbers, which would mislead "
    "anyone judging real-hardware feasibility from them alone.",
    "A named, previously-unverified hypothesis was tested correlationally, then CONFIRMED with an "
    "actual intervention: NLMS's instantaneous-power normalization measurably inflates its "
    "effective step size during fades; flooring that normalization (a new, off-by-default "
    "parameter with zero effect on any other experiment) measurably closed the Frozen-vs-DD gap in "
    "every tested cell and flipped two of four from \"DD doesn't help\" to \"DD measurably helps\" "
    "— turning a best-supported explanation into a demonstrated cause.",
    "The same discipline was applied to a second previously-open question — and this time the "
    "'confirmed' root cause turned out to be wrong. The high-SNR CNN error floor was first "
    "attributed, on strong correlational grounds (χ²=168.0, p=3.5×10⁻³⁶ positional clustering), to "
    "a window/overlap-add reconstruction artifact. A code-correctness critique pass found the real "
    "cause: reconstruct_from_windows silently zero-filled trailing samples no window covered, read "
    "by a hard-decision demodulator as a fixed, wrong-half-the-time bit — this project's SEVENTH "
    "real bug. Fixing it collapsed BPSK's floor (constant 83/100,000 errors at 10-20dB) to "
    "0-1/100,000, and QPSK's floor at 15-20dB to 0; the clustering statistic recomputes to χ²≈0.74 "
    "(p≈0.86, indistinguishable from uniform) on the corrected data — the original result was itself "
    "manufactured by the bug. A smaller, genuine residual floor survives only at QPSK 10dB, now "
    "explained by high-noise-magnitude samples rather than reconstruction position. The negative "
    "result on the triangular-weighting intervention was correct, but for a different reason than "
    "understood at the time: reweighting cannot fix a sample no window covers at all.",
    "This project's single-training-draw practice was given a first quantitative check, in one "
    "narrow regime, not validated project-wide: "
    "training-draw variance is consistently smaller than test-time Monte Carlo variance across all "
    "4 conditions tested (ratio 0.00-0.32x), for Case Study 1's CNN only.",
    "RLS, added as a third classical baseline, required fixing a real safety bug (this project's "
    "fifth) before its first use, then a first pass without tail-averaging found it measurably "
    "worse than LMS/NLMS at low-to-moderate SNR. Adding tail-averaging (the same hardening LMS/NLMS "
    "already had) and "
    "re-running — after catching and correcting an initial comparison against a stale results "
    "file — found RLS now matches or beats LMS/NLMS at EVERY SNR level tested, confirming the "
    "original theoretical motivation once given a fair comparison.",
    "A missing standard baseline (Zero-Forcing) was found by searching comparable published work, "
    "not by re-reading this project's own code — once a later bug fix (see below) corrected both "
    "genie equalizers' numbers, it confirmed the textbook prediction directly: ZF (no "
    "noise-awareness) is the slightly worse of the two genie equalizers for QPSK at moderate SNR, "
    "both comfortably beating No-Processing.",
    "A deliberate adversarial self-critique pass (four independent reviewer agents, each briefed to "
    "find real problems) surfaced a sixth bug (an asymmetric RRC filter, src/signal_gen.py, impact "
    "quantified rather than assumed small), a genuine reporting error in this project's own "
    "diagnostic (undefined-ratio denominators counted as \"not inflated\"), an uncorrected-multiple-"
    "comparisons gap in Section 4.5 (now corrected with both Bonferroni and Benjamini-Hochberg), "
    "and the preamble-drift correction above — each fixed, corrected, or explicitly caveated rather "
    "than left unaddressed.",
    "An outside-reviewer critique pass (deliberately unfamiliar with this project's internal "
    "history) found a gap none of the prior self-critique rounds had named: every CNN result trains "
    "and tests on the exact same known channel, while LMS/NLMS/MMSE/ZF all re-estimate or are "
    "handed the channel per-trial regardless. Testing this directly confirmed it is a real, "
    "substantial effect: a CNN frozen after training on one channel degrades up to 129.71x on a "
    "held-out channel with a different tap-phase structure (QPSK, 10dB), while LMS and NLMS on the "
    "identical held-out channel degrade only 1.17x and 1.41x. This revises, rather than overturns, "
    "Section 9.5's practical recommendation: the CNN's advantage is real when training and "
    "deployment channels match, but should not be assumed to survive a channel that drifts from "
    "that training distribution.",
    "A final code-correctness critique pass found an EIGHTH real bug — and this one had produced a "
    "wrong conclusion backed by a real, honest, three-part investigation, not just a correlational "
    "guess. design_mmse_equalizer computed its autocorrelation via np.correlate, which silently "
    "returns the complex conjugate of the value its own documented Wiener-Hopf formula requires — "
    "negligible for BPSK, severe for QPSK/16-QAM. This was the actual cause of both the \"QPSK genie "
    "MMSE gets worse at high SNR\" finding and the \"16-QAM genie MMSE regresses below doing "
    "nothing\" finding — both previously investigated with real phase-bias, tap-count, and "
    "frequency-response checks that were each individually correct but collectively missed the one "
    "place the bug lived. Fixed and re-verified against a from-scratch ground-truth solution; "
    "re-running Case Study 2, the severity sweep, Case Study 4, and the ZF comparison at full scale "
    "reverses both findings completely — genie MMSE now ties or beats every method at high SNR for "
    "QPSK and 16-QAM. Unlike the seventh bug (a since-explained statistical artifact), here every "
    "individual verification step taken at the time was itself sound — the lesson: a mechanistic "
    "investigation can rule out several real alternative explanations correctly and still miss the "
    "actual cause, if it never inspects the specific line of code computing the number in question.",
])

doc.add_page_break()

# ===========================================================================
# 11. Assumptions, Limitations, Future Work
# ===========================================================================
add_heading(doc, "11. Assumptions, Limitations, and Future Work", level=1)
add_bullets(doc, [
    "Single static multipath channel in Case Study 2's main run — partially addressed by the "
    "severity sweep (varies magnitude, not delay spread or tap count).",
    "LMS/NLMS use one fixed μ across the entire SNR sweep; the residual ~0.2-0.3dB SNR shortfall at "
    "high SNR in Case Study 1 is a direct, quantifiable consequence.",
    "The theoretical Q-function BER curve is valid for Case Study 1 but has been removed (not just "
    "captioned) from Case Study 2's figures, since it is not a bound there.",
    "The genie MMSE equalizer is a linear, symbol-spaced reference bound, not a universal ceiling "
    "(the fractionally-spaced caveat in Section 4.7 still applies). An earlier version of this "
    "bullet claimed the ceiling gets more consequential for 16-QAM -- that was this project's "
    "eighth real bug, corrected: the properly-computed bound stays close to optimal at every "
    "modulation tested.",
    "The decision-directed reliability gate is dead code for any SPS>1 signal — affects every "
    "experiment using pulse shaping, all of which use the (unaffected-in-practice) default "
    "enable_decision_directed=False.",
    "The boundary-aware-loss, training-variance, 16-QAM, LMS-vs-NLMS-asymmetry, and CNN-high-SNR-"
    "floor follow-ups were all run at reduced scale (4,000-30,000 symbols, 4-5 MC trials, sometimes "
    "only 5 independent draws) for speed — treat their percentages, ratios, and correlation "
    "coefficients as indicative at that reduced-confidence standard, not full precision.",
    "A sixth real bug (found via an adversarial self-critique pass): rrc_filter centered its time "
    "axis on N/2 rather than (N-1)/2, making the transmit/receive filter measurably asymmetric. "
    "Fixed and verified symmetric. Impact was QUANTIFIED, not assumed: a targeted before/after "
    "comparison (20,000-symbol BPSK, Case Study 2's channel) gave BER 61/20,000 (old) vs. 55/20,000 "
    "(new) at 0dB and identical 0 errors at 10-20dB, output SNR shifting only ~0.01dB — small and "
    "directionally favorable, but Case Studies 2/4, the severity sweep, and Case Study 3's SPS=4 "
    "portion still predate this fix and were not re-run at full scale (the RLS and ZF follow-ups "
    "were run after the fix and already reflect it; the boundary-loss and training-variance "
    "follow-ups never used this filter at all, so are unaffected regardless).",
    "The \"channel looks static during the preamble\" premise for Case Study 3 was asserted, not "
    "measured, and found false for the actual default parameters (Section 6.1) — this does not "
    "invalidate Section 6.4's Frozen-vs-DD comparison, but changes the causal story for why DD helps.",
    "Section 4.5's Fisher's-exact \"statistically real\" claims (9 tests) — DONE, corrected: "
    "Bonferroni leaves 3/9 significant, Benjamini-Hochberg FDR (the more defensible choice for this "
    "exploratory family) leaves 7/9 significant. The qualitative conclusion (a small, practically-"
    "negligible floor) is unaffected, but NLMS's specific high-SNR regression is not statistically "
    "distinguishable from chance under either correction.",
    "The reconstruction-tail-fill bug (Section 3.7's correction, this project's seventh) affects any "
    "CNN/Hybrid experiment whose signal length isn't an exact multiple of the 64-sample stride "
    "relative to the 128-sample window, not just the diagnostic that found it. Case Study 1 (100,000 "
    "symbols, SPS=1, a 32-sample uncovered tail) HAS been re-run with the fix and its tables reflect "
    "it: CNN's floor at both modulations' 15-20dB is now exactly 0 (previously 1.51e-4/1.71e-4); "
    "Hybrid's floor dropped substantially but not to zero, now matching LMS's own residual "
    "misadjustment floor. Case Study 4 (25,000 symbols, SPS=4, also a 32-sample tail) and the "
    "boundary-aware-loss/training-variance follow-ups (25,000/15,000 symbols, SPS=1) have NOT yet "
    "been re-run and still show pre-fix numbers — Case Study 2, the severity sweep, and Case Study "
    "3's SPS=4 portion use signal lengths that divide the stride exactly and were never affected "
    "regardless of when they were run.",
    "The autocorrelation-conjugate bug (Sections 4.7/4.10's correction, this project's eighth) "
    "affects every complex (non-BPSK) use of design_mmse_equalizer/design_zf_equalizer -- unlike the "
    "seventh bug above, every affected experiment HAS been re-run at full scale, with no remaining "
    "gap. Case Study 2, the severity sweep, Case Study 4's 16-QAM comparison, and the ZF comparison "
    "all use this equalizer for QPSK and/or 16-QAM and have all been re-run with the fix; their "
    "tables and prose throughout this report already reflect the corrected numbers. BPSK-only "
    "content is unaffected regardless (real symbols mostly mask a conjugate error).",
])

add_heading(doc, "Named, Scoped-Out Future Work", level=2)
add_para(doc, "Eight items from the original future-work list have since been addressed and are marked accordingly; each surfaced new, more specific follow-ups of its own.")
add_bullets(doc, [
    "DONE — RLS as a third classical baseline (Section 4.9): a first pass without tail-averaging "
    "found RLS worse than LMS/NLMS at low-to-moderate SNR. DONE — the tail-averaging fix itself: "
    "implemented and tested (not left as a guess), and found to substantially improve RLS at every "
    "SNR level — RLS now matches or beats LMS/NLMS across the board. Still open: an untuned λ=0.99 "
    "forgetting factor as a further refinement.",
    "MLSE/Viterbi equalization — the classical gold-standard for ISI channels — still open.",
    "DONE (Case Study 3) — a genuinely time-varying channel. DONE — root-cause the LMS-vs-NLMS "
    "tracking asymmetry, AND the named intervention (flooring NLMS's normalization) was actually "
    "tested and confirmed to close the gap. Still open: fix the DD update rule for SPS>1; re-run "
    "the flat-fading control at SPS=4 once that fix exists.",
    "DONE for 16-QAM (Case Study 4) — a working, self-tested 8-PSK implementation exists but "
    "wasn't run through the full pipeline; also open: root-causing 16-QAM's hard No-Processing "
    "floor more precisely (the genie MMSE \"regression\" once suspected here turned out to be "
    "this project's eighth bug, not a real effect — see Section 4.7/7's correction).",
    "DONE — root-cause the loss-independent high-SNR error floor: this was first attributed to a "
    "window/overlap-add reconstruction artifact (χ²=168.0, p=3.5×10⁻³⁶) — a conclusion later found "
    "to be wrong (this project's seventh bug, a reconstruction zero-fill error, Section 3.7). Still "
    "open, on the smaller residual that survives the fix at QPSK 10dB: whether it can be closed "
    "with a wider window or more training data.",
    "DONE — multiple independent training draws: training-draw variance is consistently smaller "
    "than test-time variance (ratio 0.00-0.32x across all 4 conditions tested) — not yet extended "
    "to Hybrid or Case Study 2's ISI channel.",
    "DONE — Zero-Forcing (ZF) as a linear-equalizer baseline, found missing by comparing to "
    "published literature rather than this project's own code: once the eighth bug was fixed, "
    "confirms the textbook prediction (ZF slightly worse than MMSE) at a realistic magnitude, "
    "rather than the bug-inflated \"regression\" first reported.",
    "Cycle-accurate hardware validation of the compute-cost analysis — the MAC-count-based "
    "throughput requirements were not validated against an actual microcontroller, DSP, or NPU.",
    "A full-scale re-run of Case Studies 2-4 and the severity sweep with the corrected (symmetric) "
    "RRC filter — a small, targeted check found the effect is small and directionally favorable, "
    "but a complete re-run to confirm this at full precision has not been done.",
])

doc.add_page_break()

# ===========================================================================
# References
# ===========================================================================
add_heading(doc, "References", level=1)
add_para(doc,
    "All findings, numbers, tables, and figures elsewhere in this report are this project's own "
    "original results. The sources below are cited by number at the specific claims they support "
    "in Section 9.5; they were used only to contextualize how this project's own measurements "
    "compare to prior published work, never as a substitute for this project's own measurements."
)
references = [
    "S. Kumar et al., \"Comparative Analysis of LMS, NLMS, and RLS Adaptive Filters in Vehicle "
    "Automation Systems under Mixed Noise Conditions,\" ResearchGate, 2025.",
    "\"Comparative Study of ZF, LMS and RLS Adaptive Equalization Techniques,\" arXiv:2312.06084, 2023.",
    "Ignitarium / Neurealm, \"Adaptive Filters for Signal Processing: A Comparative Study,\" "
    "Medium / company engineering blog, n.d.",
    "\"Channel Equalization of Adaptive Filters Using LMS and RLS Algorithms,\" International "
    "Journal of Creative Research Thoughts (IJCRT), rjpn.org, n.d.",
    "R. Krishna, \"Autoencoders vs FIR Filters: Smarter Signal Denoising with ML and DSP,\" Medium, "
    "2025.",
    "Omdena, \"Denoising Autoencoders,\" Omdena engineering blog, n.d.",
]
add_numbered(doc, references)

doc.add_page_break()

# ===========================================================================
# Appendix
# ===========================================================================
add_heading(doc, "Appendix: Project Structure & Regression Test Suite", level=1)
add_para(doc, "A CI-enforced, 68-test regression suite (tests/, ~2.5min to run, run automatically on "
    "every push via GitHub Actions) codifies every invariant discovered during this project:")
add_bullets(doc, [
    "test_lms_stability.py — LMS/NLMS never diverge below No-Processing; DD mode off by default.",
    "test_rls_stability.py — RLS never diverges below No-Processing; DD mode off by default "
    "(16 tests, added when RLS was hardened before its first real use, Section 4.9).",
    "test_receiver_frontend.py — matched filter recovers unit gain; No-Processing's SNR matches "
    "nominal (the identity that caught two of the three Case Study 2 bugs).",
    "test_channel.py — add_multipath is a true identity and causal (the identity that caught the "
    "third bug); add_time_varying_multipath is causal, matches the static channel at zero drift, "
    "and correctly fixes the direct path unless vary_direct_path=True (added during the "
    "self-critique pass — previously zero coverage on this function).",
    "test_mmse_equalizer.py — genie MMSE never loses to No-Processing (now with QPSK coverage, added after the eighth bug went undetected in an all-BPSK suite); ZF converges to MMSE as noise vanishes and recovers symbols near-perfectly at negligible noise; a ground-truth Wiener-Hopf solve directly verifies design_mmse_equalizer's output against a from-scratch reimplementation, the test that would have caught the autocorrelation-conjugate bug (Section 4.7).",
    "test_cnn_autoencoder.py — reconstruct_from_windows' uncovered-tail gap and its fallback fix are locked in with 5 tests (skipped, not failed, when TensorFlow isn't installed, matching this project's lightweight-CI convention) — previously zero coverage on this function, the gap that produced Section 3.7's seventh bug.",
    "test_no_leakage.py — no train/test bit-sequence collisions.",
    "test_metrics.py — BER/error-count consistency, rule-of-three sanity.",
])
add_para(doc,
    "Every experiment referenced in this report lives as a standalone Python script under "
    "experiments/ (e.g. run_case1_no_isi.py, run_case2_isi.py, run_severity_sweep.py, "
    "run_timevarying_channel.py, run_16qam_comparison.py, run_case1_boundary_loss.py, "
    "compute_cost_analysis.py, run_rls_comparison.py, run_zf_comparison.py, "
    "diagnose_lms_nlms_asymmetry.py, diagnose_cnn_high_snr_floor.py, diagnose_training_variance.py, "
    "test_nlms_floor_intervention.py, test_cnn_overlap_weighting_intervention.py, "
    "test_cnn_channel_generalization.py), with shared "
    "logic in src/ (signal_gen.py, channel.py, lms_filter.py, nlms_filter.py, rls_filter.py, "
    "mmse_equalizer.py, utils.py, metrics.py) and a CNN-specific module (cnn_boundary_aware.py). "
    "Raw results are stored as CSVs under results/tables/ and figures under results/figures/, "
    "organized by case study. This report is the complete, self-contained writeup of that "
    "codebase's methodology, results, and conclusions."
)

set_update_fields_on_open(doc)
doc.save(str(OUT_PATH))
print(f"Saved: {OUT_PATH}")
